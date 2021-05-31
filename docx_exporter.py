from docx import Document
from docx.shared import Pt

def export_problem_to_docx(problems: list, name: str = 'problem'):
    problem_document = Document()
    table = problem_document.add_table(rows=0, cols=2)
    table.autofit = True
    problems_count = len(problems)
    odd = True
    if problems_count % 2 == 0:
        odd = False
    now = 0
    while now < problems_count:
        row_cells = table.add_row().cells
        row_cells[0].text = str(now + 1) + ') ' + problems[now][0]
        now += 1
        if now == problems_count:
            if odd:
                break
        row_cells[1].text = str(now + 1) + ') ' + problems[now][0]
        now += 1

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(16)

    problem_document.save(name + '.docx')

def export_solution_to_docx(problems: list, name: str = 'solution'):
    solution_document = Document()
    table = solution_document.add_table(rows=0, cols=2)
    table.autofit = True
    problems_count = len(problems)
    odd = True
    if problems_count % 2 == 0:
        odd = False
    now = 0
    while now < problems_count:
        row_cells = table.add_row().cells
        row_cells[0].text = str(now + 1) + ') ' + str(problems[now][1])
        now += 1
        if now == problems_count:
            if odd:
                break
        row_cells[1].text = str(now + 1) + ') ' + str(problems[now][1])
        now += 1

    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size= Pt(16)

    solution_document.save(name + '.docx')
